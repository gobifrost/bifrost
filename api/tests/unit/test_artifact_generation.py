import io

from docx import Document
from openpyxl import load_workbook

from shared.artifact_generation import (
    generate_document,
    generate_document_with_images,
    generate_spreadsheet,
    generate_text,
    validate_artifact_content,
)
from src.models.contracts.artifacts import (
    DocumentArtifactSpec,
    SpreadsheetArtifactSpec,
    TextArtifactSpec,
)


def _png() -> bytes:
    from PIL import Image

    buffer = io.BytesIO()
    Image.new("RGB", (80, 40), "navy").save(buffer, format="PNG")
    return buffer.getvalue()


def _document_spec(output_format: str) -> DocumentArtifactSpec:
    return DocumentArtifactSpec.model_validate(
        {
            "filename": "customer-brief",
            "format": output_format,
            "title": "Customer brief",
            "subtitle": "Decision-ready summary",
            "sections": [
                {
                    "heading": "Decision",
                    "paragraphs": ["Proceed with the rollout."],
                    "bullets": ["Owner assigned", "Review scheduled"],
                    "table": {
                        "columns": ["Phase", "Owner"],
                        "rows": [["Pilot", "Operations"]],
                    },
                }
            ],
        }
    )


def test_generate_pdf_uses_flowing_document_contract() -> None:
    artifact = generate_document(_document_spec("pdf"))

    assert artifact.filename == "Customer Brief.pdf"
    assert artifact.content_type == "application/pdf"
    assert artifact.content.startswith(b"%PDF-")
    assert artifact.size_bytes > 500


def test_generate_docx_can_be_opened_by_python_docx() -> None:
    artifact = generate_document(_document_spec("docx"))
    document = Document(io.BytesIO(artifact.content))

    assert artifact.filename == "Customer Brief.docx"
    assert document.paragraphs[0].text == "Customer brief"
    assert any(paragraph.text == "Proceed with the rollout." for paragraph in document.paragraphs)
    assert document.tables[0].cell(1, 0).text == "Pilot"


def test_generate_docx_embeds_image_from_workspace_path() -> None:
    spec = DocumentArtifactSpec.model_validate(
        {
            "filename": "field-report",
            "format": "docx",
            "title": "Field report",
            "sections": [
                {
                    "heading": "Photo",
                    "images": [
                        {"path": "Bluetick Portrait.png", "caption": "Bluetick"}
                    ],
                }
            ],
        }
    )

    artifact = generate_document_with_images(
        spec,
        {"Bluetick Portrait.png": _png()},
    )
    document = Document(io.BytesIO(artifact.content))

    assert len(document.inline_shapes) == 1
    assert any(paragraph.text == "Bluetick" for paragraph in document.paragraphs)


def test_generate_xlsx_styles_and_validates_workbook() -> None:
    spec = SpreadsheetArtifactSpec.model_validate(
        {
            "filename": "adoption-dashboard",
            "sheets": [
                {
                    "name": "Adoption",
                    "columns": ["Channel", "Status"],
                    "rows": [["Partner", "Ready"], ["Direct", "Pilot"]],
                }
            ],
        }
    )

    artifact = generate_spreadsheet(spec)
    workbook = load_workbook(io.BytesIO(artifact.content))
    worksheet = workbook["Adoption"]

    assert artifact.filename == "Adoption Dashboard.xlsx"
    assert worksheet["A2"].value == "Partner"
    assert worksheet.freeze_panes == "A2"
    assert worksheet.auto_filter.ref == "A1:B3"
    assert worksheet["A1"].font.bold is True
    workbook.close()


def test_generate_json_normalizes_valid_content() -> None:
    artifact = generate_text(
        TextArtifactSpec(
            filename="result",
            format="json",
            content='{"ready":true,"count":2}',
        )
    )

    assert artifact.filename == "Result.json"


def test_generate_markdown_uses_human_friendly_name_and_extension() -> None:
    artifact = generate_text(
        TextArtifactSpec(
            filename="customer-health_report.md",
            format="markdown",
            content="# Health report",
        )
    )

    assert artifact.filename == "Customer Health Report.md"
    assert artifact.content == b"# Health report"


def test_video_artifact_validation_accepts_mp4_signature() -> None:
    validate_artifact_content(
        filename="Demo.mp4",
        content_type="video/mp4",
        content=b"\x00\x00\x00\x18ftypmp42\x00\x00\x00\x00mp42isom",
    )
