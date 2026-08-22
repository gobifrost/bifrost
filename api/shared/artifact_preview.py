"""Safe HTML previews for generated Office artifacts."""

from __future__ import annotations

import io
from html import escape

from docx import Document
from openpyxl import load_workbook

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

_STYLE = """
<style>
  :root { color-scheme: light dark; font-family: ui-sans-serif, system-ui, sans-serif; }
  body { margin: 0; padding: 28px; color: #172033; background: #fff; line-height: 1.55; }
  main { max-width: 900px; margin: 0 auto; }
  h1, h2, h3 { line-height: 1.2; margin: 1.5em 0 .55em; }
  h1:first-child { margin-top: 0; }
  p { margin: .65em 0; }
  table { width: 100%; border-collapse: collapse; margin: 20px 0 30px; font-size: 13px; }
  th, td { border: 1px solid #d5dbea; padding: 7px 9px; text-align: left; vertical-align: top; }
  th { background: #edf1f7; font-weight: 650; }
  .sheet { margin-bottom: 44px; overflow-x: auto; }
  .truncated { color: #67728a; font-size: 12px; }
  @media (prefers-color-scheme: dark) {
    body { color: #e6eaf2; background: #11141a; }
    th, td { border-color: #343b49; }
    th { background: #202631; }
    .truncated { color: #9aa5ba; }
  }
</style>
"""


def _page(body: str) -> str:
    return f"<!doctype html><html><head><meta charset='utf-8'>{_STYLE}</head><body><main>{body}</main></body></html>"


def preview_docx(content: bytes) -> str:
    """Render paragraphs and tables from a DOCX without executing active content."""
    document = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs[:500]:
        text = escape(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("title"):
            parts.append(f"<h1>{text}</h1>")
        elif style.startswith("heading 1"):
            parts.append(f"<h2>{text}</h2>")
        elif style.startswith("heading"):
            parts.append(f"<h3>{text}</h3>")
        elif "list" in style:
            parts.append(f"<p>• {text}</p>")
        else:
            parts.append(f"<p>{text}</p>")
    for table in document.tables[:30]:
        rows: list[str] = []
        for row_index, row in enumerate(table.rows[:200]):
            tag = "th" if row_index == 0 else "td"
            cells = "".join(f"<{tag}>{escape(cell.text)}</{tag}>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        parts.append(f"<table>{''.join(rows)}</table>")
    return _page("".join(parts) or "<p>No previewable document content.</p>")


def preview_xlsx(content: bytes) -> str:
    """Render a bounded workbook preview without evaluating formulas."""
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=False)
    parts: list[str] = []
    try:
        for worksheet in workbook.worksheets[:5]:
            rows: list[str] = []
            for row_index, row in enumerate(worksheet.iter_rows(max_row=201, values_only=True)):
                tag = "th" if row_index == 0 else "td"
                cells = "".join(
                    f"<{tag}>{escape('' if value is None else str(value))}</{tag}>"
                    for value in row[:100]
                )
                rows.append(f"<tr>{cells}</tr>")
            parts.append(
                f"<section class='sheet'><h2>{escape(worksheet.title)}</h2>"
                f"<table>{''.join(rows)}</table>"
                + (
                    "<p class='truncated'>Preview limited to the first 200 rows.</p>"
                    if worksheet.max_row > 201
                    else ""
                )
                + "</section>"
            )
    finally:
        workbook.close()
    return _page("".join(parts) or "<p>No previewable workbook content.</p>")


def preview_office_artifact(content: bytes, content_type: str) -> str | None:
    if content_type == DOCX_CONTENT_TYPE:
        return preview_docx(content)
    if content_type == XLSX_CONTENT_TYPE:
        return preview_xlsx(content)
    return None
