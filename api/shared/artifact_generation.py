"""Trusted, schema-first generators for Chat and workflow artifacts.

The model supplies bounded structured content. This module owns layout and
binary serialization, so no provider needs to emit binary data or execute
arbitrary Python.
"""

from __future__ import annotations

import io
import json
import zipfile
from dataclasses import dataclass
from decimal import Decimal
from html import escape
from pathlib import Path
from collections.abc import Mapping
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as ReportLabImage,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src.models.contracts.artifacts import (
    DocumentArtifactSpec,
    SpreadsheetArtifactSpec,
    TextArtifactSpec,
)

MAX_ARTIFACT_BYTES = 25 * 1024 * 1024
MAX_VIDEO_ARTIFACT_BYTES = 250 * 1024 * 1024

_CONTENT_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv",
    "html": "text/html",
    "markdown": "text/markdown",
    "text": "text/plain",
    "json": "application/json",
}


class ArtifactGenerationError(ValueError):
    """Raised when a structured artifact payload cannot produce a safe file."""


@dataclass(frozen=True)
class GeneratedArtifact:
    filename: str
    content_type: str
    content: bytes
    provider: str | None = None
    model: str | None = None
    input_tokens: int = 0
    output_tokens: int = 0
    provider_cost: Decimal | None = None

    @property
    def size_bytes(self) -> int:
        return len(self.content)


def validate_artifact_content(
    *, filename: str, content_type: str, content: bytes
) -> None:
    """Validate signatures for trusted generators and promoted workflow files."""
    if not content:
        raise ArtifactGenerationError(f"{filename} is empty.")
    max_bytes = (
        MAX_VIDEO_ARTIFACT_BYTES
        if content_type in {"video/mp4", "video/webm"}
        else MAX_ARTIFACT_BYTES
    )
    if len(content) > max_bytes:
        raise ArtifactGenerationError(
            f"{filename} exceeds the {max_bytes // (1024 * 1024)} MB limit."
        )
    if content_type == "application/pdf":
        if not content.startswith(b"%PDF-"):
            raise ArtifactGenerationError(f"{filename} is not a valid PDF.")
        return
    if content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }:
        expected = (
            "word/document.xml"
            if "wordprocessingml" in content_type
            else "xl/workbook.xml"
        )
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                if expected not in archive.namelist():
                    raise ArtifactGenerationError(
                        f"{filename} does not match its declared Office format."
                    )
        except zipfile.BadZipFile as exc:
            raise ArtifactGenerationError(
                f"{filename} is not a valid Office document."
            ) from exc
        return
    if content_type in {"image/png", "image/jpeg", "image/webp", "image/gif"}:
        expected_format = {
            "image/png": "PNG",
            "image/jpeg": "JPEG",
            "image/webp": "WEBP",
            "image/gif": "GIF",
        }[content_type]
        try:
            with Image.open(io.BytesIO(content)) as image:
                image.verify()
                if image.format != expected_format:
                    raise ArtifactGenerationError(
                        f"{filename} does not match its declared image type."
                    )
        except (OSError, ValueError) as exc:
            raise ArtifactGenerationError(f"{filename} is not a valid image.") from exc
        return
    if content_type == "video/mp4":
        if len(content) < 12 or content[4:8] != b"ftyp":
            raise ArtifactGenerationError(f"{filename} is not a valid MP4 video.")
        return
    if content_type == "video/webm":
        if not content.startswith(b"\x1a\x45\xdf\xa3"):
            raise ArtifactGenerationError(f"{filename} is not a valid WebM video.")
        return
    if content_type.startswith("text/") or content_type in {
        "application/json",
        "application/csv",
    }:
        try:
            decoded = content.decode("utf-8")
            if content_type == "application/json":
                json.loads(decoded)
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise ArtifactGenerationError(
                f"{filename} does not match its declared text format."
            ) from exc
        return
    raise ArtifactGenerationError(
        f"{filename} uses unsupported artifact type {content_type}."
    )


def safe_artifact_filename(filename: str, extension: str) -> str:
    name = Path(filename.replace("\\", "/")).name.strip()
    if not name or name in {".", ".."}:
        raise ArtifactGenerationError("Artifact filename is invalid.")
    public_extension = {"markdown": "md", "text": "txt"}.get(extension, extension)
    expected = f".{public_extension}"
    stem = name[: -len(expected)] if name.lower().endswith(expected) else name
    words = stem.replace("_", " ").replace("-", " ").split()
    proper_stem = " ".join(
        word.upper()
        if any(character.isdigit() for character in word)
        else f"{word[:1].upper()}{word[1:]}"
        for word in words
    )
    return f"{proper_stem}{expected}"


def _finish(filename: str, output_format: str, content: bytes) -> GeneratedArtifact:
    if not content:
        raise ArtifactGenerationError("Artifact generation produced an empty file.")
    if len(content) > MAX_ARTIFACT_BYTES:
        raise ArtifactGenerationError("Generated artifact exceeds the 25 MB limit.")
    artifact = GeneratedArtifact(
        filename=safe_artifact_filename(filename, output_format),
        content_type=_CONTENT_TYPES[output_format],
        content=content,
    )
    validate_artifact_content(
        filename=artifact.filename,
        content_type=artifact.content_type,
        content=artifact.content,
    )
    return artifact


def generate_document(spec: DocumentArtifactSpec) -> GeneratedArtifact:
    """Generate a PDF or DOCX from a flowing document specification."""
    if spec.format == "pdf":
        return _generate_pdf(spec)
    return _generate_docx(spec)


def generate_document_with_images(
    spec: DocumentArtifactSpec,
    image_content: Mapping[str, bytes],
) -> GeneratedArtifact:
    """Generate a document after resolving its opaque image references."""
    if spec.format == "pdf":
        return _generate_pdf(spec, image_content)
    return _generate_docx(spec, image_content)


def _resolved_document_image(
    path: str,
    image_content: Mapping[str, bytes],
) -> bytes:
    content = image_content.get(path)
    if content is None:
        raise ArtifactGenerationError(
            f"Document image {path} was not found in the artifact workspace."
        )
    return content


def _generate_pdf(
    spec: DocumentArtifactSpec,
    image_content: Mapping[str, bytes] | None = None,
) -> GeneratedArtifact:
    buffer = io.BytesIO()
    page_size = LETTER if spec.page_size == "letter" else A4
    document = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=spec.title,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ArtifactTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        spaceAfter=16,
    )
    subtitle_style = ParagraphStyle(
        "ArtifactSubtitle",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#475569"),
        spaceAfter=24,
    )
    story: list[Any] = [Paragraph(escape(spec.title), title_style)]
    if spec.subtitle:
        story.append(Paragraph(escape(spec.subtitle), subtitle_style))

    for index, section in enumerate(spec.sections):
        if index and section.heading:
            story.append(Spacer(1, 10))
        if section.heading:
            story.append(Paragraph(escape(section.heading), styles["Heading2"]))
        for paragraph in section.paragraphs:
            story.append(Paragraph(escape(paragraph).replace("\n", "<br/>"), styles["BodyText"]))
            story.append(Spacer(1, 7))
        if section.bullets:
            story.append(
                ListFlowable(
                    [
                        ListItem(Paragraph(escape(item), styles["BodyText"]))
                        for item in section.bullets
                    ],
                    bulletType="bullet",
                    leftIndent=18,
                )
            )
            story.append(Spacer(1, 9))
        if section.table:
            rows = [section.table.columns, *section.table.rows]
            table_data = [
                [Paragraph(escape("" if cell is None else str(cell)), styles["BodyText"]) for cell in row]
                for row in rows
            ]
            table = Table(table_data, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.extend([table, Spacer(1, 10)])
        for image_spec in section.images:
            content = _resolved_document_image(
                image_spec.path,
                image_content or {},
            )
            with Image.open(io.BytesIO(content)) as source:
                width_px, height_px = source.size
            max_width = min(image_spec.max_width_inches * inch, document.width)
            rendered_height = max_width * height_px / width_px
            story.append(
                ReportLabImage(
                    io.BytesIO(content),
                    width=max_width,
                    height=rendered_height,
                    hAlign="CENTER",
                )
            )
            if image_spec.caption:
                story.append(
                    Paragraph(
                        escape(image_spec.caption),
                        ParagraphStyle(
                            "ArtifactImageCaption",
                            parent=styles["Caption"],
                            alignment=TA_CENTER,
                            textColor=colors.HexColor("#64748B"),
                            spaceBefore=5,
                        ),
                    )
                )
            story.append(Spacer(1, 10))

    def add_page_number(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawRightString(page_size[0] - 0.75 * inch, 0.45 * inch, f"Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    content = buffer.getvalue()
    if not content.startswith(b"%PDF-"):
        raise ArtifactGenerationError("PDF validation failed.")
    return _finish(spec.filename, "pdf", content)


def _generate_docx(
    spec: DocumentArtifactSpec,
    image_content: Mapping[str, bytes] | None = None,
) -> GeneratedArtifact:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_heading(spec.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if spec.subtitle:
        subtitle = document.add_paragraph(spec.subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(11)
            run.font.italic = True

    for item in spec.sections:
        if item.heading:
            document.add_heading(item.heading, level=1)
        for paragraph in item.paragraphs:
            document.add_paragraph(paragraph)
        for bullet in item.bullets:
            document.add_paragraph(bullet, style="List Bullet")
        if item.table:
            table = document.add_table(rows=1, cols=len(item.table.columns))
            table.style = "Table Grid"
            for index, column in enumerate(item.table.columns):
                table.rows[0].cells[index].text = column
            for row in item.table.rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = "" if value is None else str(value)
        for image_spec in item.images:
            content = _resolved_document_image(
                image_spec.path,
                image_content or {},
            )
            image_paragraph = document.add_paragraph()
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.add_run().add_picture(
                io.BytesIO(content),
                width=Inches(image_spec.max_width_inches),
            )
            if image_spec.caption:
                caption = document.add_paragraph(image_spec.caption)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    content = buffer.getvalue()
    if not content.startswith(b"PK"):
        raise ArtifactGenerationError("DOCX validation failed.")
    return _finish(spec.filename, "docx", content)


def generate_spreadsheet(spec: SpreadsheetArtifactSpec) -> GeneratedArtifact:
    """Generate a styled XLSX workbook from bounded tabular data."""
    workbook = Workbook()
    workbook.remove(workbook.active)
    header_fill = PatternFill("solid", fgColor="1E3A5F")
    header_font = Font(color="FFFFFF", bold=True)

    for sheet_spec in spec.sheets:
        worksheet = workbook.create_sheet(title=sheet_spec.name)
        worksheet.append(sheet_spec.columns)
        for row in sheet_spec.rows:
            worksheet.append(row)
        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(vertical="center")
        worksheet.row_dimensions[1].height = 22
        if sheet_spec.freeze_header:
            worksheet.freeze_panes = "A2"
        if sheet_spec.auto_filter:
            worksheet.auto_filter.ref = worksheet.dimensions
        for column_index, column_name in enumerate(sheet_spec.columns, start=1):
            values = [column_name, *(row[column_index - 1] for row in sheet_spec.rows[:500])]
            width = min(max(len(str(value)) if value is not None else 0 for value in values) + 2, 60)
            worksheet.column_dimensions[get_column_letter(column_index)].width = max(width, 10)

    buffer = io.BytesIO()
    workbook.save(buffer)
    content = buffer.getvalue()
    try:
        checked = load_workbook(io.BytesIO(content), read_only=True, data_only=False)
        checked.close()
    except Exception as exc:
        raise ArtifactGenerationError("XLSX validation failed.") from exc
    return _finish(spec.filename, "xlsx", content)


def generate_text(spec: TextArtifactSpec) -> GeneratedArtifact:
    """Generate a UTF-8 text-family artifact with lightweight validation."""
    content = spec.content
    if spec.format == "json":
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as exc:
            raise ArtifactGenerationError("JSON artifacts must contain valid JSON.") from exc
        content = json.dumps(parsed, ensure_ascii=False, indent=2)
    return _finish(spec.filename, spec.format, content.encode("utf-8"))
